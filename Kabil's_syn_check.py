
# This module is totally independent of changes needed in additional files. Variable part is separated bellow, so don't worry.

import os

from docx import Document
# from docx.shared import Pt
# from docx.oxml import OxmlElement
# from docx.enum.text import WD_UNDERLINE



doc_path = r"C:\Users\WALTON\OneDrive\Desktop\Tanvir\Vocabulary\1.docx" # .....................................................................

def kabil_syn_search_capital():
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            word_serial_no = row.cells[0].text
            target_word_cell = row.cells[6]  # Assuming the 2nd cell contains the target word...........................................................
            # target_word = target_word_cell.text.strip()
            target_words = target_word_cell.text.strip().split()
            
            for target_word in target_words:
            
                print(f'The target word is- {word_serial_no}.{target_word}')

                for cell in row.cells[4:6]:  # Iterate through cells in the same row, starting from 2 to 2..............................................
                            # row.cells[start:end] will include the cell at index 'start' but not the cell at index 'end'.
                    for paragraph in cell.paragraphs:
                        updated_text = []  # Create an empty list to store updated words
                        for run in paragraph.runs:
                            words = run.text.split()  # Split the text into words.
                            for i, word in enumerate(words):
                                
                                if target_word.lower() in word.lower():
                                    
                                    # If the word matches, uppercase it
                                    words[i] = word.upper()
                                    # words[i] = word.upper() assigns the uppercase version of word back to the list of words (words) at the same index i where word was found.
                                    
                            updated_text.extend(words)  # Extend the list with updated words
                        # Combine words back into a paragraph
                        paragraph.text = ' '.join(updated_text) #
            # After this loop, 'paragraph.text' will contain the updated text with the specified words in uppercase.

    doc.save(doc_path)

    # os.system(f'start {doc_path}')


# kabil_syn_search_capital()







def kabil_syn_not_matched():
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            word_serial_no = row.cells[0].text
            target_word_cell = row.cells[6]  # Assuming the 2nd cell contains the target word...........................................................
            # target_word = target_word_cell.text.strip()
            target_words = target_word_cell.text.strip().split()
            
            
            updated_kabil = []
            
            for target_word in target_words:
                
                x = 0

                print(f'The target word is- {word_serial_no}.{target_word}')

                for cell in row.cells[4:6]:  # Iterate through cells in the same row, starting from 2 to 2..............................................
                            # row.cells[start:end] will include the cell at index 'start' but not the cell at index 'end'.
                    for paragraph in cell.paragraphs:
                        updated_text = []  # Create an empty list to store updated words
                        for run in paragraph.runs:
                            words = run.text.split()  # Split the text into words.
                            for i, word in enumerate(words):
                                
                                if target_word.lower() in word.lower():
                                    
                                    # If the word matches, uppercase it
                                    words[i] = word.upper()
                                    # words[i] = word.upper() assigns the uppercase version of word back to the list of words (words) at the same index i where word was found.
                                    
                                    x += 1
                                    
                            updated_text.extend(words)  # Extend the list with updated words
                        # Combine words back into a paragraph
                        paragraph.text = ' '.join(updated_text) #
            # After this loop, 'paragraph.text' will contain the updated text with the specified words in uppercase.
                    
                if x > 0:
                    target_word = ""
                    updated_kabil.extend([target_word])
                else:
                    updated_kabil.extend([target_word])
                    
            list_to_text = ', '.join(updated_kabil)

            print(f"The outpurt result after search is---------- {list_to_text}.")
            
            target_word_cell.text = list_to_text

    doc.save(doc_path)



# kabil_syn_not_matched()